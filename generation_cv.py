#!/usr/bin/env python
# -*- coding: utf-8 -*-

def generate_cv(json_path, output_filename):
    import json
    import os
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Funzione per creare una linea orizzontale nel documento
    def add_horizontal_line(paragraph):
        p = paragraph._p  # p è l'elemento xml che rappresenta il paragrafo
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        
        # Aggiungi la linea inferiore
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Spessore della linea
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')  # Colore blu
        pBdr.append(bottom)

    # Funzione per applicare lo stile ai titoli
    def style_heading(paragraph, text, size, bold=True, color=RGBColor(68, 114, 196)):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        run.bold = bold
        font = run.font
        font.size = Pt(size)
        font.color.rgb = color

    # Carica i dati dal file JSON
    with open(json_path, 'r', encoding='utf-8') as file:
        cv_data = json.load(file)

    # Crea un nuovo documento Word
    doc = Document()

    # Configura i margini del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Crea una tabella per l'intestazione con foto (se presente)
    has_photo = 'photo' in cv_data['basics'] and cv_data['basics']['photo']
    if has_photo:
        # Usa tabella per mettere foto a destra e nome/tagline a sinistra
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        
        # Colonna sinistra per nome e tagline
        left_cell = header_table.cell(0, 0)
        left_cell.width = Inches(5.5)
        
        # Nome e titolo
        name_para = left_cell.paragraphs[0]
        name_run = name_para.add_run(cv_data['basics']['name'])
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.color.rgb = RGBColor(68, 114, 196)  # Blu moderno
        
        # Tagline sotto il nome
        tagline_para = left_cell.add_paragraph()
        tagline_run = tagline_para.add_run(cv_data['basics']['tagline'])
        tagline_run.italic = True
        tagline_run.font.size = Pt(12)
        
        # Colonna destra per la foto
        right_cell = header_table.cell(0, 1)
        right_cell.width = Inches(1.5)
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Aggiungi la foto
        photo_path = os.path.join('static', cv_data['basics']['photo'])
        if os.path.exists(photo_path):
            try:
                photo_run = right_cell.paragraphs[0].add_run()
                photo_run.add_picture(photo_path, width=Inches(1.3), height=Inches(1.3))
            except Exception as e:
                print(f"Errore nel caricare l'immagine: {e}")
    else:
        # Intestazione standard senza foto
        header = doc.add_paragraph()
        name_run = header.add_run(cv_data['basics']['name'])
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.color.rgb = RGBColor(68, 114, 196)  # Blu moderno

        # Tagline sotto il nome
        tagline_para = doc.add_paragraph()
        tagline_run = tagline_para.add_run(cv_data['basics']['tagline'])
        tagline_run.italic = True
        tagline_run.font.size = Pt(12)

    # Linea separatrice
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    # Informazioni di contatto
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Email
    email = contact_para.add_run(f"📧 {cv_data['basics']['email']}  ")
    email.font.size = Pt(10)

    # Telefono
    phone_mobile = cv_data['basics']['phone']['mobile']
    phone = contact_para.add_run(f"📱 {phone_mobile}  ")
    phone.font.size = Pt(10)

    # GitHub
    if 'profiles' in cv_data['basics'] and 'github' in cv_data['basics']['profiles']:
        github = contact_para.add_run(f"GitHub: {cv_data['basics']['profiles']['github']}")
        github.font.size = Pt(10)

    # Indirizzo e nazionalità
    location_para = doc.add_paragraph()
    location = location_para.add_run(f"🏠 {cv_data['basics']['location']}  •  🌍 Nazionalità: {cv_data['basics']['nationality']}")
    location.font.size = Pt(10)

    # Data e luogo di nascita
    birth_para = doc.add_paragraph()
    birth_info = birth_para.add_run(f"📅 Nato il: {cv_data['basics']['birth']['date']} a {cv_data['basics']['birth']['place']}")
    birth_info.font.size = Pt(10)

    doc.add_paragraph()  # Spazio

    # Esperienza lavorativa
    style_heading(doc.add_paragraph(), "ESPERIENZA LAVORATIVA", 16)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    for job in cv_data['work']:
        # Titolo del lavoro e azienda
        job_title = doc.add_paragraph()
        position_run = job_title.add_run(f"{job['position']} presso {job['company']}")
        position_run.bold = True
        position_run.font.size = Pt(12)
        position_run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Durata
        if job['duration']:
            duration_para = doc.add_paragraph()
            duration_run = duration_para.add_run(f"⏱️ {job['duration']}")
            duration_run.italic = True
            duration_run.font.size = Pt(10)
        
        # Achievements come elenco puntato
        for achievement in job['achievements']:
            achievement_para = doc.add_paragraph(style='List Bullet')
            achievement_run = achievement_para.add_run(achievement)
            achievement_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spazio tra lavori

    # Istruzione
    style_heading(doc.add_paragraph(), "ISTRUZIONE E FORMAZIONE", 16)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    for edu in cv_data['education']:
        # Titolo dell'istruzione
        edu_para = doc.add_paragraph()
        degree_run = edu_para.add_run(f"{edu['degree']}")
        degree_run.bold = True
        degree_run.font.size = Pt(12)
        degree_run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Istituzione
        inst_para = doc.add_paragraph()
        institution_run = inst_para.add_run(f"📚 {edu['institution']}")
        institution_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spazio tra istruzioni

    # Competenze tecniche
    style_heading(doc.add_paragraph(), "COMPETENZE", 16)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    # Competenze AI
    if 'ai' in cv_data['skills']:
        skill_para = doc.add_paragraph()
        skill_title = skill_para.add_run("🤖 AI & Machine Learning: ")
        skill_title.bold = True
        skill_title.font.size = Pt(11)
        skill_content = skill_para.add_run(", ".join(cv_data['skills']['ai']))
        skill_content.font.size = Pt(10)

    # Programmazione
    if 'programming' in cv_data['skills']:
        prog_para = doc.add_paragraph()
        prog_title = prog_para.add_run("💻 Programmazione: ")
        prog_title.bold = True
        prog_title.font.size = Pt(11)
        
        # Avanzato
        if 'advanced' in cv_data['skills']['programming']:
            advanced = prog_para.add_run("\nAvanzato: ")
            advanced.bold = True
            advanced.font.size = Pt(10)
            adv_content = prog_para.add_run(", ".join(cv_data['skills']['programming']['advanced']))
            adv_content.font.size = Pt(10)
        
        # Intermedio
        if 'intermediate' in cv_data['skills']['programming']:
            intermediate = prog_para.add_run("\nIntermedio: ")
            intermediate.bold = True
            intermediate.font.size = Pt(10)
            int_content = prog_para.add_run(", ".join(cv_data['skills']['programming']['intermediate']))
            int_content.font.size = Pt(10)
        
        # Base
        if 'basic' in cv_data['skills']['programming']:
            basic = prog_para.add_run("\nBase: ")
            basic.bold = True
            basic.font.size = Pt(10)
            basic_content = prog_para.add_run(", ".join(cv_data['skills']['programming']['basic']))
            basic_content.font.size = Pt(10)

    # Automazione industriale
    if 'industrialAutomation' in cv_data['skills']:
        indust_para = doc.add_paragraph()
        indust_title = indust_para.add_run("🏭 Automazione Industriale: ")
        indust_title.bold = True
        indust_title.font.size = Pt(11)
        indust_content = indust_para.add_run(", ".join(cv_data['skills']['industrialAutomation']))
        indust_content.font.size = Pt(10)

    # Sistemi operativi
    if 'systems' in cv_data['skills']:
        sys_para = doc.add_paragraph()
        sys_title = sys_para.add_run("💽 Sistemi Operativi: ")
        sys_title.bold = True
        sys_title.font.size = Pt(11)
        
        if 'windows' in cv_data['skills']['systems']:
            windows = sys_para.add_run(f"\nWindows: {cv_data['skills']['systems']['windows']}")
            windows.font.size = Pt(10)
        
        if 'linux' in cv_data['skills']['systems']:
            linux = sys_para.add_run(f"\nLinux: {cv_data['skills']['systems']['linux']}")
            linux.font.size = Pt(10)

    # Software
    if 'software' in cv_data['skills']:
        sw_para = doc.add_paragraph()
        sw_title = sw_para.add_run("🖥️ Software: ")
        sw_title.bold = True
        sw_title.font.size = Pt(11)
        sw_content = sw_para.add_run(", ".join(cv_data['skills']['software']))
        sw_content.font.size = Pt(10)

    # DevOps
    if 'devOps' in cv_data['skills']:
        devops_para = doc.add_paragraph()
        devops_title = devops_para.add_run("🔄 DevOps & Database: ")
        devops_title.bold = True
        devops_title.font.size = Pt(11)
        devops_content = devops_para.add_run(", ".join(cv_data['skills']['devOps']))
        devops_content.font.size = Pt(10)

    doc.add_paragraph()  # Spazio

    # Lingue
    style_heading(doc.add_paragraph(), "LINGUE", 16)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    for lang in cv_data['languages']:
        lang_para = doc.add_paragraph()
        lang_name = lang_para.add_run(f"{lang['language']}: ")
        lang_name.bold = True
        lang_name.font.size = Pt(11)
        lang_level = lang_para.add_run(f"{lang['level']}")
        lang_level.font.size = Pt(10)

    doc.add_paragraph()  # Spazio

    # Altre informazioni
    style_heading(doc.add_paragraph(), "ALTRE INFORMAZIONI", 16)
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)

    # Patente
    if 'drivingLicense' in cv_data['other']:
        license_para = doc.add_paragraph()
        license_title = license_para.add_run("🚗 Patente: ")
        license_title.bold = True
        license_title.font.size = Pt(11)
        license_content = license_para.add_run(cv_data['other']['drivingLicense'])
        license_content.font.size = Pt(10)

    # Hobby
    if 'hobbies' in cv_data['other']:
        hobby_para = doc.add_paragraph()
        hobby_title = hobby_para.add_run("🎮 Hobby: ")
        hobby_title.bold = True
        hobby_title.font.size = Pt(11)
        for idx, hobby in enumerate(cv_data['other']['hobbies']):
            hobby_content = hobby_para.add_run(f"{hobby}")
            hobby_content.font.size = Pt(10)
            if idx < len(cv_data['other']['hobbies']) - 1:
                hobby_para.add_run(", ")

    # Qualità
    if 'qualities' in cv_data['other']:
        quality_para = doc.add_paragraph()
        quality_title = quality_para.add_run("✨ Qualità personali: ")
        quality_title.bold = True
        quality_title.font.size = Pt(11)
        quality_content = quality_para.add_run(", ".join(cv_data['other']['qualities']))
        quality_content.font.size = Pt(10)

    # Competenze digitali
    if 'digitalSkills' in cv_data:
        doc.add_paragraph()  # Spazio
        style_heading(doc.add_paragraph(), "COMPETENZE DIGITALI", 16)
        line_para = doc.add_paragraph()
        add_horizontal_line(line_para)
        
        for digital_skill in cv_data['digitalSkills']:
            skill_para = doc.add_paragraph()
            skill_name = skill_para.add_run(f"{digital_skill['skill']}: ")
            skill_name.bold = True
            skill_name.font.size = Pt(11)
            skill_level = skill_para.add_run(f"{digital_skill['level']}")
            skill_level.font.size = Pt(10)

    # Piè di pagina
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer_para.add_run("CV generato automaticamente - " + cv_data['basics']['name'])
    footer_text.font.size = Pt(8)
    footer_text.font.color.rgb = RGBColor(128, 128, 128)

    # Salva il documento
    doc.save(output_filename)
    print(f"CV creato con successo: {output_filename}")